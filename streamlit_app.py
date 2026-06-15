# streamlit_app.py
# Aplicación Streamlit para diseño de muros estructurales
# Basado en el código de diseño de muro con As1 calculado con acero real del alma.

import io
import tempfile
from pathlib import Path

import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


from dataclasses import dataclass, asdict
from math import sqrt, ceil, pi
from typing import Dict, Any, Optional

import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, Circle


# ============================================================
# FUNCIONES BÁSICAS
# ============================================================

def area_barra_cm2(db_mm: float) -> float:
    """Área de una varilla circular en cm²."""
    db_cm = db_mm / 10.0
    return pi * db_cm**2 / 4.0


def n_barras_requeridas(as_req_cm2: float, db_mm: float) -> int:
    """Número entero de barras para cubrir un área de acero requerida."""
    if as_req_cm2 <= 1e-9:
        return 0
    return max(1, ceil(as_req_cm2 / area_barra_cm2(db_mm)))


# ============================================================
# DATOS DE ENTRADA
# ============================================================

@dataclass
class DatosMuro:
    # Geometría
    lw_m: float = 2.40
    bw_m: float = 0.30
    hw_m: float = 33.66
    n_pisos: int = 11

    # Materiales
    fc_kgcm2: float = 280.0
    fy_kgcm2: float = 4200.0
    fyt_kgcm2: float = 4200.0

    # Solicitaciones
    Pu_tonf: float = 352.0
    Vu_tonf: float = 60.0
    Mu_plano_tonfm: float = 208.0   # Si es 0, diseña como compresión axial predominante.

    # Factores usados en el ejemplo del documento
    phi_v: float = 0.75
    phi_flexion: float = 0.90
    phi_o: float = 1.40

    # Cuantías mínimas distribuidas
    rho_vertical: float = 0.0025
    rho_horizontal: float = 0.0025

    # Desplazamiento máximo de techo para condición de borde.
    # El programa calcula internamente δu/hw = desplazamiento_techo / hw.
    desplazamiento_techo_cm: float = 14.96

    # Armado colocado por el usuario
    db_borde_mm: float = 14.0

    # Armado longitudinal colocado en CADA cabezal/elemento de borde.
    # La sección del cabezal se dibuja como un rectángulo lbe x bw.
    # Siempre se colocan 4 barras en las esquinas.
    # n_barras_intermedias_x: barras adicionales en cada cara larga X del cabezal.
    #   Total aportado por X = 2*n_barras_intermedias_x
    # n_barras_intermedias_y: barras adicionales en cada cara corta Y del cabezal.
    #   Total aportado por Y = 2*n_barras_intermedias_y
    n_barras_esquinas: int = 4
    n_barras_intermedias_x: int = 2
    n_barras_intermedias_y: int = 0

    db_estribo_mm: float = 10.0
    db_web_mm: float = 14.0
    s_web_cm: float = 25.0

    # Parámetros del equilibrio usado en el ejemplo
    xp_factor_lw: float = 0.095

    # Parámetros para Ash
    bc_ash_1_cm: float = 40.0
    bc_ash_2_cm: float = 26.7

    # Recubrimiento libre del elemento de borde.
    # Para calcular hx automático, el programa obtiene internamente la distancia
    # al centro de la barra extrema:
    # rec_libre + db_estribo + db_longitudinal/2.
    recubrimiento_libre_cm: float = 4.0

    # Modo de cálculo de hx:
    # - "auto": calcula hx internamente.
    # - "manual": usa hx_manual_mm ingresado por el usuario.
    modo_hx: str = "auto"
    hx_manual_mm: float = 200.0

    # Compresión axial predominante
    k_compresion: float = 1.0
    altura_libre_m: Optional[float] = None
    phi_compresion: float = 0.65


# ============================================================
# DISEÑO DEL MURO
# ============================================================

class DisenoMuro:
    def __init__(self, datos: DatosMuro):
        self.d = datos

    @property
    def lw_cm(self):
        return self.d.lw_m * 100.0

    @property
    def bw_cm(self):
        return self.d.bw_m * 100.0

    @property
    def hw_cm(self):
        return self.d.hw_m * 100.0

    @property
    def Ag_cm2(self):
        return self.lw_cm * self.bw_cm

    def modo_diseno(self) -> str:
        if abs(self.d.Mu_plano_tonfm) < 1e-9:
            return "COMPRESION_AXIAL_PREDOMINANTE"
        return "FLEXO_COMPRESION_EN_EL_PLANO"

    # --------------------------------------------------------
    # 1. CORTANTE EN EL PLANO
    # --------------------------------------------------------
    def diseno_cortante(self) -> Dict[str, Any]:
        omega = round(1.3 + self.d.n_pisos / 30.0, 2)
        Vu_amplificado = omega * self.d.phi_o * self.d.Vu_tonf

        phiVn = (
            self.d.phi_v
            * 1.6
            * sqrt(self.d.fc_kgcm2)
            * self.bw_cm
            * self.lw_cm
            / 1000.0
        )

        lw_req = (
            Vu_amplificado * 1000.0
            / (self.d.phi_v * 1.6 * sqrt(self.d.fc_kgcm2) * self.bw_cm)
        )

        bw_req = (
            Vu_amplificado * 1000.0
            / (self.d.phi_v * 1.6 * sqrt(self.d.fc_kgcm2) * self.lw_cm)
        )

        return {
            "omega": omega,
            "Vu_amplificado_tonf": Vu_amplificado,
            "phiVn_tonf": phiVn,
            "lw_req_cm": lw_req,
            "bw_req_cm": bw_req,
            "cumple_cortante": phiVn >= Vu_amplificado,
        }

    # --------------------------------------------------------
    # 2A. COMPRESIÓN AXIAL PREDOMINANTE
    # --------------------------------------------------------
    def diseno_compresion_axial(self) -> Dict[str, Any]:
        lc_m = self.d.altura_libre_m if self.d.altura_libre_m is not None else self.d.hw_m
        h_cm = self.bw_cm

        slender = self.d.k_compresion * (lc_m * 100.0) / h_cm
        factor_esbeltez = max(1.0 - (slender / 32.0) ** 2, 0.0)

        phiPn = (
            self.d.phi_compresion
            * 0.55
            * self.d.fc_kgcm2
            * self.Ag_cm2
            * factor_esbeltez
            / 1000.0
        )

        h_min = max(10.0, self.lw_cm / 25.0)

        return {
            "modo": "COMPRESION_AXIAL_PREDOMINANTE",
            "k_lc_h": slender,
            "factor_esbeltez": factor_esbeltez,
            "phiPn_tonf": phiPn,
            "h_min_cm": h_min,
            "cumple_espesor": self.bw_cm >= h_min,
            "cumple_axial": phiPn >= self.d.Pu_tonf,
        }

    # --------------------------------------------------------
    # 2B. FLEXO-COMPRESIÓN EN EL PLANO
    # --------------------------------------------------------
    def diseno_flexocompresion(self) -> Dict[str, Any]:
        if self.modo_diseno() == "COMPRESION_AXIAL_PREDOMINANTE":
            return self.diseno_compresion_axial()

        lw = self.d.lw_m

        Mn = self.d.Mu_plano_tonfm / self.d.phi_flexion
        xp_m = self.d.xp_factor_lw * lw

        # Acero distribuido vertical REAL colocado en el alma que aporta como Ts1.
        # Se considera el acero vertical en dos caras dentro del tramo 0.60*lw,
        # siguiendo la misma zona usada por el equilibrio del ejemplo.
        area_web = area_barra_cm2(self.d.db_web_mm)
        As_vertical_por_metro = 2.0 * area_web * 100.0 / self.d.s_web_cm
        longitud_alma_efectiva_m = 0.60 * lw
        As1 = As_vertical_por_metro * longitud_alma_efectiva_m
        Ts1 = As1 * self.d.fy_kgcm2 / 1000.0

        brazo_Ts1 = 0.40 * lw
        brazo_Ts2 = 0.80 * lw

        Ts2 = (Mn - self.d.Pu_tonf * xp_m - Ts1 * brazo_Ts1) / brazo_Ts2
        Ts2 = max(Ts2, 0.0)

        As2_req = Ts2 * 1000.0 / self.d.fy_kgcm2

        n_requeridas = n_barras_requeridas(As2_req, self.d.db_borde_mm)

        # Barras colocadas por el usuario en CADA cabezal.
        # Mínimo físico: 4 barras en las esquinas.
        n_esquinas = max(4, int(self.d.n_barras_esquinas))
        n_x = max(0, int(self.d.n_barras_intermedias_x))
        n_y = max(0, int(self.d.n_barras_intermedias_y))

        n_colocadas = n_esquinas + 2 * n_x + 2 * n_y
        armado_es_propuesto = False

        As2_col = n_colocadas * area_barra_cm2(self.d.db_borde_mm)
        cumple_borde = As2_col >= As2_req
        As2_faltante = max(As2_req - As2_col, 0.0)
        As2_excedente = max(As2_col - As2_req, 0.0)
        requiere_acero_adicional_borde_por_flexion = As2_req > 1e-9

        return {
            "modo": "FLEXO_COMPRESION_EN_EL_PLANO",
            "Mn_tonfm": Mn,
            "xp_m": xp_m,
            "As1_cm2": As1,
            "As_vertical_alma_por_metro_cm2_m": As_vertical_por_metro,
            "longitud_alma_efectiva_m": longitud_alma_efectiva_m,
            "Ts1_tonf": Ts1,
            "Ts2_tonf": Ts2,
            "As2_req_cm2": As2_req,
            "db_borde_mm": self.d.db_borde_mm,
            "n_barras_borde_requeridas": n_requeridas,
            "n_barras_borde_colocadas": n_colocadas,
            "n_barras_esquinas": n_esquinas,
            "n_barras_intermedias_x": n_x,
            "n_barras_intermedias_y": n_y,
            "As2_colocado_cm2": As2_col,
            "As2_faltante_cm2": As2_faltante,
            "As2_excedente_cm2": As2_excedente,
            "requiere_acero_adicional_borde_por_flexion": requiere_acero_adicional_borde_por_flexion,
            "cumple_acero_borde": cumple_borde,
            "armado_borde": (
                f"{n_colocadas} Ø{self.d.db_borde_mm:.0f} mm "
                f"(4 esquinas + {n_x} intermedias en X por cara + {n_y} intermedias en Y por cara)"
            ),
            "armado_borde_requerido_minimo": f"{n_requeridas} Ø{self.d.db_borde_mm:.0f} mm",
            "armado_es_propuesto_por_programa": armado_es_propuesto,
        }

    # --------------------------------------------------------
    # 3. ELEMENTOS ESPECIALES DE BORDE
    # --------------------------------------------------------
    def elementos_borde(self) -> Dict[str, Any]:
        if self.modo_diseno() == "COMPRESION_AXIAL_PREDOMINANTE":
            return {
                "requiere_borde_especial": False,
                "motivo": "No se revisan elementos especiales de borde porque Mu_plano = 0. El diseño principal es por compresión axial predominante.",
            }

        pu_agfc = self.d.Pu_tonf * 1000.0 / (self.Ag_cm2 * self.d.fc_kgcm2)

        delta_u_sobre_hw = (self.d.desplazamiento_techo_cm / self.hw_cm) if self.hw_cm > 0 else 0.0
        denom = 600.0 * (1.5 * delta_u_sobre_hw)
        c_cm = self.lw_cm / denom if denom > 0 else float("inf")
        c_sobre_lw = c_cm / self.lw_cm

        lbe_cm = max(c_cm - 0.10 * self.lw_cm, c_cm / 2.0)

        # En este código se conserva el criterio del ejemplo del documento:
        # si hay flexo-compresión en el plano, se calcula y muestra el elemento especial de borde.
        # El usuario debe interpretar esto como necesidad de confinamiento del borde comprimido,
        # no necesariamente como demanda de acero adicional de tracción.
        requiere_borde_especial = True

        return {
            "requiere_borde_especial": requiere_borde_especial,
            "Pu_sobre_Agfc": pu_agfc,
            "desplazamiento_techo_cm": self.d.desplazamiento_techo_cm,
            "delta_u_sobre_hw": delta_u_sobre_hw,
            "c_cm": c_cm,
            "c_sobre_lw": c_sobre_lw,
            "lbe_cm": lbe_cm,
            "motivo": "Se revisan elementos especiales de borde porque existe flexo-compresión en el plano del muro.",
        }


    def calcular_hx_mm(self, lbe_cm: float) -> Dict[str, Any]:
        """
        Calcula hx automáticamente o toma hx manual según el modo elegido.

        Concepto:
        hx es la mayor separación centro a centro entre barras longitudinales
        consecutivas que están soportadas lateralmente por estribos, vinchas o ganchos.

        Si el detalle real tiene más vinchas/ramales que el modelo automático,
        conviene usar modo manual e ingresar el hx real medido del detalle.
        """
        rec_libre_cm = max(0.0, self.d.recubrimiento_libre_cm)
        db_estribo_cm = self.d.db_estribo_mm / 10.0
        db_long_cm = self.d.db_borde_mm / 10.0

        # Distancia desde el borde exterior hasta el centro de la barra extrema.
        # Se calcula internamente a partir del recubrimiento puro.
        rec_cm = rec_libre_cm + db_estribo_cm + db_long_cm / 2.0

        x_libre_cm = max(lbe_cm - 2.0 * rec_cm, 0.0)
        y_libre_cm = max(self.bw_cm - 2.0 * rec_cm, 0.0)

        n_x = max(0, int(self.d.n_barras_intermedias_x))
        n_y = max(0, int(self.d.n_barras_intermedias_y))

        # Separación máxima entre barras en las caras largas X.
        sx_cm = x_libre_cm / (n_x + 1) if (n_x + 1) > 0 else x_libre_cm

        # Separación máxima entre barras en las caras cortas Y.
        sy_cm = y_libre_cm / (n_y + 1) if (n_y + 1) > 0 else y_libre_cm

        hx_auto_cm = max(sx_cm, sy_cm)
        hx_auto_mm = hx_auto_cm * 10.0

        if str(self.d.modo_hx).lower().startswith("manual"):
            hx_mm = max(1.0, float(self.d.hx_manual_mm))
            hx_cm = hx_mm / 10.0
            modo_usado = "manual"
        else:
            hx_mm = hx_auto_mm
            hx_cm = hx_auto_cm
            modo_usado = "automático"

        return {
            "modo_hx": modo_usado,
            "recubrimiento_libre_cm": rec_libre_cm,
            "distancia_centro_barra_extrema_cm": rec_cm,
            "x_libre_cm": x_libre_cm,
            "y_libre_cm": y_libre_cm,
            "sx_cm": sx_cm,
            "sy_cm": sy_cm,
            "hx_auto_cm": hx_auto_cm,
            "hx_auto_mm": hx_auto_mm,
            "hx_cm": hx_cm,
            "hx_mm": hx_mm,
        }


    # --------------------------------------------------------
    # 4. DETALLAMIENTO DE BORDE
    # --------------------------------------------------------
    def detallamiento_borde(self) -> Dict[str, Any]:
        borde = self.elementos_borde()

        if not borde["requiere_borde_especial"]:
            return borde

        db_long_cm = self.d.db_borde_mm / 10.0

        # Longitud de anclaje / gancho usada para reproducir el ejemplo.
        ldh = max(
            8.0 * db_long_cm,
            15.0,
            0.076 * self.d.fy_kgcm2 * db_long_cm / sqrt(self.d.fc_kgcm2),
        )

        # Espaciamiento de estribos. hx se calcula automáticamente desde el detalle de barras.
        hx_info = self.calcular_hx_mm(borde["lbe_cm"])
        hx_mm = hx_info["hx_mm"]

        s1 = self.bw_cm / 4.0
        s2 = 6.0 * db_long_cm

        # Expresión tipo ACI/NEC para s0.
        # IMPORTANTE: s0 no debe hacerse negativo cuando hx es grande.
        # Usamos la limitación usual: s0 no mayor que 150 mm y no menor que 100 mm.
        s0_mm_raw = 100.0 + (350.0 - hx_mm) / 3.0
        s0_mm = min(150.0, max(100.0, s0_mm_raw))
        s0 = s0_mm / 10.0

        candidatos_s = {
            "bw/4": s1,
            "6db longitudinal": s2,
            "s0(hx)": s0,
        }
        s = min(candidatos_s.values())
        controla_s = min(candidatos_s, key=candidatos_s.get)

        # Seguridad numérica: no permitir separación negativa o cero.
        s = max(s, 1e-6)

        # En el ejemplo gobierna 0.09*s*bc*fc/fyt.
        Ash1_req = max(0.09 * s * self.d.bc_ash_1_cm * self.d.fc_kgcm2 / self.d.fyt_kgcm2, 0.0)
        Ash2_req = max(0.09 * s * self.d.bc_ash_2_cm * self.d.fc_kgcm2 / self.d.fyt_kgcm2, 0.0)

        area_estribo = area_barra_cm2(self.d.db_estribo_mm)

        n_ramas_1 = n_barras_requeridas(Ash1_req, self.d.db_estribo_mm)
        n_ramas_2 = n_barras_requeridas(Ash2_req, self.d.db_estribo_mm)

        # Detalle esquemático concatenado con los ramales requeridos.
        # Un estribo cerrado aporta 2 ramas en una dirección.
        # Si se requieren más de 2 ramas, se agregan vinchas interiores.
        n_ramas_control = max(n_ramas_1, n_ramas_2)
        n_vinchas = max(0, n_ramas_control - 2)
        if n_vinchas == 0:
            detalle_estribos = f"Estribo cerrado Ø{self.d.db_estribo_mm:.0f} @ {s:.1f} cm"
        elif n_vinchas == 1:
            detalle_estribos = f"Estribo cerrado Ø{self.d.db_estribo_mm:.0f} + 1 vincha Ø{self.d.db_estribo_mm:.0f} @ {s:.1f} cm"
        else:
            detalle_estribos = f"Estribo cerrado Ø{self.d.db_estribo_mm:.0f} + {n_vinchas} vinchas Ø{self.d.db_estribo_mm:.0f} @ {s:.1f} cm"

        return {
            **borde,
            "ldh_cm": ldh,
            "s_estribos_cm": s,
            "s1_bw_4_cm": s1,
            "s2_6db_cm": s2,
            "s0_raw_mm": s0_mm_raw,
            "s0_usado_cm": s0,
            "controla_s": controla_s,
            "modo_hx": hx_info["modo_hx"],
            "hx_mm": hx_mm,
            "hx_cm": hx_info["hx_cm"],
            "hx_auto_mm": hx_info["hx_auto_mm"],
            "hx_auto_cm": hx_info["hx_auto_cm"],
            "hx_sx_cm": hx_info["sx_cm"],
            "hx_sy_cm": hx_info["sy_cm"],
            "recubrimiento_libre_cm": hx_info["recubrimiento_libre_cm"],
            "distancia_centro_barra_extrema_cm": hx_info["distancia_centro_barra_extrema_cm"],
            "Ash1_req_cm2": Ash1_req,
            "Ash2_req_cm2": Ash2_req,
            "n_ramas_Ash1": n_ramas_1,
            "n_ramas_Ash2": n_ramas_2,
            "Ash1_colocado_cm2": n_ramas_1 * area_estribo,
            "Ash2_colocado_cm2": n_ramas_2 * area_estribo,
            "n_ramas_control": n_ramas_control,
            "n_vinchas": n_vinchas,
            "detalle_estribos": detalle_estribos,
        }

    # --------------------------------------------------------
    # 5. ACERO DISTRIBUIDO DEL ALMA
    # --------------------------------------------------------
    def acero_distribuido(self) -> Dict[str, Any]:
        area_web = area_barra_cm2(self.d.db_web_mm)

        # Dos caras
        As_por_metro = 2.0 * area_web * 100.0 / self.d.s_web_cm
        rho_colocada = As_por_metro / (self.bw_cm * 100.0)

        return {
            "db_web_mm": self.d.db_web_mm,
            "s_web_cm": self.d.s_web_cm,
            "As_por_metro_cm2_m": As_por_metro,
            "rho_colocada": rho_colocada,
            "cumple_rho_vertical": rho_colocada >= self.d.rho_vertical,
            "cumple_rho_horizontal": rho_colocada >= self.d.rho_horizontal,
            "detalle": f"Ø{self.d.db_web_mm:.0f} mm @ {self.d.s_web_cm:.0f} cm en dos caras",
        }

    # --------------------------------------------------------
    # 6. RESULTADO GENERAL
    # --------------------------------------------------------
    def resolver(self) -> Dict[str, Any]:
        return {
            "entrada": asdict(self.d),
            "modo": self.modo_diseno(),
            "geometria": {
                "lw_cm": self.lw_cm,
                "bw_cm": self.bw_cm,
                "hw_cm": self.hw_cm,
                "Ag_cm2": self.Ag_cm2,
            },
            "cortante": self.diseno_cortante(),
            "flexocompresion_o_compresion": self.diseno_flexocompresion(),
            "borde": self.detallamiento_borde(),
            "acero_distribuido": self.acero_distribuido(),
        }

    # --------------------------------------------------------
    # 7. VALIDACIÓN CONTRA EL EJEMPLO 128-137
    # --------------------------------------------------------
    def validar_ejemplo_documento(self, tol_rel: float = 0.03) -> Dict[str, Any]:
        """
        Valida solo cuando los datos corresponden al ejemplo del documento.
        Si se cambian los datos, la validación puede salir REVISAR porque ya no es el mismo ejemplo.
        """
        r = self.resolver()

        if self.modo_diseno() != "FLEXO_COMPRESION_EN_EL_PLANO":
            return {
                "aplica_validacion": False,
                "ok_validacion": None,
                "mensaje": "La validación del ejemplo 128-137 aplica al caso con flexo-compresión en el plano. Para Mu_plano=0 no aplica esa comparación.",
                "comparacion": {},
            }

        targets = {
            "omega": 1.67,
            "Vu_amplificado_tonf": 140.3,
            "lw_req_cm": 233.0,
            "bw_req_cm": 29.0,
            "Mn_tonfm": 231.0,
            "Ts1_tonf": 45.36,
            "Ts2_tonf": 55.9,
            "As2_req_cm2": 13.3,
            "Pu_sobre_Agfc": 0.17,
            "c_sobre_lw": 0.25,
            "c_cm": 60.0,
            "lbe_cm": 36.0,
            "ldh_cm": 26.6,
            "s_estribos_cm": 7.5,
            "Ash1_req_cm2": 1.8,
            "Ash2_req_cm2": 1.2,
        }

        flat = {}
        for sec in ["cortante", "flexocompresion_o_compresion", "borde"]:
            for k, v in r[sec].items():
                flat[k] = v

        comparacion = {}
        ok_global = True

        for k, ref in targets.items():
            calc = flat.get(k)
            if calc is None:
                err = None
                ok = False
            else:
                err = abs(calc - ref) / abs(ref)
                ok = err <= tol_rel

            comparacion[k] = {
                "calculado": calc,
                "referencia": ref,
                "error_relativo": err,
                "ok": ok,
            }
            ok_global = ok_global and ok

        return {
            "aplica_validacion": True,
            "ok_validacion": ok_global,
            "tolerancia_relativa": tol_rel,
            "mensaje": "OK" if ok_global else "REVISAR: algún resultado no coincide con la tolerancia.",
            "comparacion": comparacion,
        }

    # --------------------------------------------------------
    # 8. MEMORIA EN MARKDOWN
    # --------------------------------------------------------
    def memoria_markdown(self) -> str:
        r = self.resolver()
        g = r["geometria"]
        c = r["cortante"]
        fx = r["flexocompresion_o_compresion"]
        b = r["borde"]
        web = r["acero_distribuido"]
        val = self.validar_ejemplo_documento()

        lines = []

        lines.append("# Memoria de cálculo - Muro estructural")
        lines.append("")
        lines.append("## 1. Datos")
        lines.append(f"- Longitud del muro: **lw = {self.d.lw_m:.2f} m = {g['lw_cm']:.1f} cm**")
        lines.append(f"- Espesor del muro: **bw = {self.d.bw_m:.2f} m = {g['bw_cm']:.1f} cm**")
        lines.append(f"- Altura total: **hw = {self.d.hw_m:.2f} m**")
        lines.append(f"- Área bruta: **Ag = lw·bw = {g['Ag_cm2']:.1f} cm²**")
        lines.append(f"- Hormigón: **f'c = {self.d.fc_kgcm2:.0f} kg/cm²**")
        lines.append(f"- Acero: **fy = {self.d.fy_kgcm2:.0f} kg/cm²**")
        lines.append(f"- Carga axial: **Pu = {self.d.Pu_tonf:.2f} tonf**")
        lines.append(f"- Cortante: **Vu = {self.d.Vu_tonf:.2f} tonf**")
        lines.append(f"- Momento en el plano: **Mu_plano = {self.d.Mu_plano_tonfm:.2f} t·m**")
        lines.append("")

        lines.append("## 2. Modo de diseño")
        if self.modo_diseno() == "COMPRESION_AXIAL_PREDOMINANTE":
            lines.append("Como **Mu_plano = 0**, el muro se diseña como **compresión axial predominante**.")
            lines.append("No se activan elementos especiales de borde.")
        else:
            lines.append("Como **Mu_plano ≠ 0**, el muro se diseña como **flexo-compresión en el plano**.")
            lines.append("Se revisan elementos especiales de borde.")
        lines.append("")

        lines.append("## 3. Diseño por cortante en el plano")
        lines.append(f"- **ω = 1.3 + N/30 = {c['omega']:.2f}**")
        lines.append(f"- **Vu' = ω·φo·Vu = {c['Vu_amplificado_tonf']:.2f} tonf**")
        lines.append(f"- **φVn = {c['phiVn_tonf']:.2f} tonf**")
        lines.append(f"- **lw requerido = {c['lw_req_cm']:.1f} cm**")
        lines.append(f"- **bw requerido = {c['bw_req_cm']:.1f} cm**")
        lines.append(f"- Veredicto: **{'CUMPLE' if c['cumple_cortante'] else 'NO CUMPLE'}**")
        lines.append("")

        lines.append("## 4. Compresión axial o flexo-compresión")
        if fx["modo"] == "COMPRESION_AXIAL_PREDOMINANTE":
            lines.append(f"- **k·lc/h = {fx['k_lc_h']:.2f}**")
            lines.append(f"- Factor de esbeltez = **{fx['factor_esbeltez']:.4f}**")
            lines.append(f"- **φPn = {fx['phiPn_tonf']:.2f} tonf**")
            lines.append(f"- Espesor mínimo: **hmin = {fx['h_min_cm']:.1f} cm**")
            lines.append(f"- Chequeo de espesor: **{'CUMPLE' if fx['cumple_espesor'] else 'NO CUMPLE'}**")
            lines.append(f"- Chequeo axial: **{'CUMPLE' if fx['cumple_axial'] else 'NO CUMPLE'}**")
        else:
            lines.append(f"- **Mn = Mu/φ = {fx['Mn_tonfm']:.2f} t·m**")
            lines.append(f"- **xp = {fx['xp_m']:.3f} m**")
            lines.append(f"- Acero vertical real del alma: **Ø{self.d.db_web_mm:.0f} @ {self.d.s_web_cm:.0f} cm en dos caras**")
            lines.append(f"- As vertical del alma por metro: **{fx['As_vertical_alma_por_metro_cm2_m']:.2f} cm²/m**")
            lines.append(f"- Longitud efectiva considerada del alma: **0.60·lw = {fx['longitud_alma_efectiva_m']:.2f} m**")
            lines.append(f"- **As1 real aportante = {fx['As1_cm2']:.2f} cm²**")
            lines.append(f"- **Ts1 = As1·fy = {fx['Ts1_tonf']:.2f} tonf**")
            lines.append(f"- **Ts2 = {fx['Ts2_tonf']:.2f} tonf**")
            lines.append(f"- **As2 adicional requerido por flexión = {fx['As2_req_cm2']:.2f} cm²**")

            if fx["requiere_acero_adicional_borde_por_flexion"]:
                lines.append(f"- Armado mínimo requerido por flexión con Ø{self.d.db_borde_mm:.0f}: **{fx['armado_borde_requerido_minimo']}**")
            else:
                lines.append("- **No se requiere acero adicional de borde por flexión**; el aporte del acero distribuido del alma cubre la tracción del equilibrio simplificado.")

            lines.append(f"- Armado colocado por el usuario en el cabezal: **{fx['armado_borde']}**")
            lines.append(f"- Distribución en el cabezal: **{fx['n_barras_esquinas']} barras en esquinas + {fx['n_barras_intermedias_x']} intermedias en X por cara + {fx['n_barras_intermedias_y']} intermedias en Y por cara**")
            lines.append(f"- Número total de barras por cabezal: **{fx['n_barras_borde_colocadas']} barras**")
            lines.append(f"- As colocado en el cabezal: **{fx['As2_colocado_cm2']:.2f} cm²**")

            if fx["requiere_acero_adicional_borde_por_flexion"]:
                lines.append(f"- Veredicto acero de borde por flexión: **{'CUMPLE' if fx['cumple_acero_borde'] else 'NO CUMPLE'}**")
                if not fx["cumple_acero_borde"]:
                    lines.append(f"- Falta acero en borde por flexión: **{fx['As2_faltante_cm2']:.2f} cm²**")
                else:
                    lines.append(f"- Excedente de acero en borde respecto a flexión: **{fx['As2_excedente_cm2']:.2f} cm²**")
            else:
                lines.append("- Veredicto acero adicional de borde por flexión: **NO REQUERIDO**")
                lines.append("- El acero del cabezal colocado se interpreta como **armado de detallamiento/confinamiento**, no como acero exigido por déficit de flexión.")
        lines.append("")

        lines.append("## 5. Elementos especiales de borde / cabezales")
        if b["requiere_borde_especial"]:
            lines.append("- ¿Requiere elemento especial de borde?: **SÍ**")
            lines.append("- Interpretación: el cabezal se exige por **confinamiento del borde comprimido**, aunque el acero adicional por flexión pueda salir igual a cero.")
            lines.append(f"- Desplazamiento máximo de techo: **δu = {b['desplazamiento_techo_cm']:.2f} cm**")
            lines.append(f"- Relación calculada: **δu/hw = {b['delta_u_sobre_hw']:.6f}**")
            lines.append(f"- **Pu/(Ag·f'c) = {b['Pu_sobre_Agfc']:.3f}**")
            lines.append(f"- **c = {b['c_cm']:.1f} cm**")
            lines.append(f"- **c/lw = {b['c_sobre_lw']:.3f}**")
            lines.append(f"- **lbe = {b['lbe_cm']:.1f} cm**")
            if b['lbe_cm'] > self.lw_cm:
                lines.append("- Advertencia: **lbe calculado es mayor que lw**; en la práctica se debe confinar todo el muro o revisar el desplazamiento δu usado.")
            elif b['lbe_cm'] > self.lw_cm / 2:
                lines.append("- Advertencia: **los elementos de borde se traslapan**; en la práctica puede requerirse confinar prácticamente todo el muro.")
            lines.append(f"- **ldh = {b['ldh_cm']:.1f} cm**")
            lines.append(f"- Modo de hx: **{b['modo_hx']}**")
            lines.append(f"- Recubrimiento libre ingresado: **{b['recubrimiento_libre_cm']:.1f} cm**")
            lines.append(f"- Distancia calculada al centro de barra extrema: **{b['distancia_centro_barra_extrema_cm']:.1f} cm**")
            lines.append(f"- Separación automática en X: **{b['hx_sx_cm']:.1f} cm**")
            lines.append(f"- Separación automática en Y: **{b['hx_sy_cm']:.1f} cm**")
            lines.append(f"- hx automático estimado: **{b['hx_auto_mm']:.0f} mm**")
            lines.append(f"- **hx usado para calcular s y Ash = {b['hx_mm']:.0f} mm**")
            lines.append(f"- Límite s1 = bw/4: **{b['s1_bw_4_cm']:.1f} cm**")
            lines.append(f"- Límite s2 = 6db longitudinal: **{b['s2_6db_cm']:.1f} cm**")
            lines.append(f"- s0 bruto por hx: **{b['s0_raw_mm']:.0f} mm**")
            lines.append(f"- s0 usado, limitado entre 100 y 150 mm: **{b['s0_usado_cm']:.1f} cm**")
            lines.append(f"- Controla la separación: **{b['controla_s']}**")
            lines.append(f"- **s estribos = {b['s_estribos_cm']:.1f} cm**")
            lines.append(f"- **Ash1 = {b['Ash1_req_cm2']:.2f} cm² → {b['n_ramas_Ash1']} ramas Ø{self.d.db_estribo_mm:.0f}**")
            lines.append(f"- **Ash2 = {b['Ash2_req_cm2']:.2f} cm² → {b['n_ramas_Ash2']} ramas Ø{self.d.db_estribo_mm:.0f}**")
            lines.append(f"- Ramales que controlan el detalle: **{b['n_ramas_control']}**")
            lines.append(f"- Vinchas interiores requeridas en el esquema: **{b['n_vinchas']}**")
            lines.append(f"- Detalle de confinamiento: **{b['detalle_estribos']}**")
        else:
            lines.append("- ¿Requiere elemento especial de borde?: **NO**")
            lines.append("- No se requieren elementos especiales de borde porque no hay momento en el plano del muro.")
        lines.append("")

        lines.append("## 6. Acero distribuido del alma")
        lines.append(f"- Armado adoptado: **{web['detalle']}**")
        lines.append(f"- As por metro en dos caras: **{web['As_por_metro_cm2_m']:.2f} cm²/m**")
        lines.append(f"- Cuantía colocada: **ρ = {web['rho_colocada']:.5f}**")
        lines.append(f"- Cuantía vertical mínima: **{'CUMPLE' if web['cumple_rho_vertical'] else 'NO CUMPLE'}**")
        lines.append(f"- Cuantía horizontal mínima: **{'CUMPLE' if web['cumple_rho_horizontal'] else 'NO CUMPLE'}**")
        lines.append("")

        lines.append("## 7. Validación")
        if val["aplica_validacion"]:
            lines.append(f"- Validación contra el ejemplo de páginas 128–137: **{'OK' if val['ok_validacion'] else 'REVISAR'}**")
            lines.append(f"- Tolerancia relativa usada: **{val['tolerancia_relativa']:.0%}**")
            lines.append("- Nota: ahora **As1 se calcula con el acero real colocado en el alma**. Por eso, si el armado del alma no coincide con el supuesto del ejemplo, algunos valores del ejemplo pueden cambiar.")
        else:
            lines.append(f"- {val['mensaje']}")

        return "\n".join(lines)

    # --------------------------------------------------------
    # 9. GRÁFICOS
    # --------------------------------------------------------
    def plot_corte_longitudinal(self, ax=None):
        r = self.resolver()
        b = r["borde"]

        if ax is None:
            fig, ax = plt.subplots(figsize=(10, 6))
        else:
            fig = ax.figure

        # Escala vertical didáctica
        altura_dibujo = max(8.0, self.d.n_pisos * 0.9)
        escala = altura_dibujo / self.d.hw_m

        lw = self.d.lw_m
        hw_draw = self.d.hw_m * escala

        ax.add_patch(Rectangle((0, 0), lw, hw_draw, fill=False, linewidth=2.0))

        # Bordes si aplican
        if b["requiere_borde_especial"]:
            lbe = b["lbe_cm"] / 100.0

            ax.add_patch(Rectangle((0, 0), lbe, hw_draw, alpha=0.18))
            ax.add_patch(Rectangle((lw - lbe, 0), lbe, hw_draw, alpha=0.18))

            ax.plot([lbe, lbe], [0, hw_draw], linestyle="--", linewidth=1.0)
            ax.plot([lw - lbe, lw - lbe], [0, hw_draw], linestyle="--", linewidth=1.0)

            ax.text(lbe / 2, hw_draw + 0.18, "Borde\nconfinado", ha="center", fontsize=9)
            ax.text(lw - lbe / 2, hw_draw + 0.18, "Borde\nconfinado", ha="center", fontsize=9)

        # Acero vertical distribuido
        s = self.d.s_web_cm / 100.0
        x = s / 2
        while x < lw:
            ax.plot([x, x], [0.10, hw_draw - 0.10], linewidth=0.8)
            x += s

        # Acero horizontal distribuido esquemático
        paso_y = hw_draw / max(self.d.n_pisos, 8)
        y = paso_y
        while y < hw_draw:
            ax.plot([0.03, lw - 0.03], [y, y], linewidth=0.6)
            y += paso_y

        ax.annotate("", xy=(0, -0.35), xytext=(lw, -0.35), arrowprops=dict(arrowstyle="<->"))
        ax.text(lw / 2, -0.48, f"lw = {lw:.2f} m", ha="center", va="top")

        ax.annotate("", xy=(lw + 0.20, 0), xytext=(lw + 0.20, hw_draw), arrowprops=dict(arrowstyle="<->"))
        ax.text(lw + 0.28, hw_draw / 2, f"hw = {self.d.hw_m:.2f} m", rotation=90, va="center")

        titulo = "Corte longitudinal del muro"
        if b["requiere_borde_especial"]:
            titulo += " con elementos de borde"
        else:
            titulo += " sin elementos de borde"

        ax.set_title(titulo)
        ax.text(0.02, -0.78, f"Alma: Ø{self.d.db_web_mm:.0f} @ {self.d.s_web_cm:.0f} cm en dos caras", fontsize=9, ha="left")

        if b["requiere_borde_especial"]:
            ax.text(0.02, -0.98, f"Bordes: {r['flexocompresion_o_compresion']['armado_borde']} + {b['detalle_estribos']}", fontsize=9, ha="left")

        ax.set_aspect("equal")
        ax.set_xlim(-0.15, lw + 0.70)
        ax.set_ylim(-1.15, hw_draw + 0.80)
        ax.axis("off")
        return fig, ax


    def coordenadas_barras_cabezal(self, x0: float, lbe: float, bw: float, rec: float = 0.055):
        """
        Coordenadas de barras dentro de un cabezal en vista transversal.
        x0: coordenada inicial del cabezal.
        lbe: longitud del cabezal en m.
        bw: espesor del muro en m.

        Distribución:
        - 4 esquinas.
        - n_barras_intermedias_x en cada cara X: cara inferior y superior.
        - n_barras_intermedias_y en cada cara Y: cara izquierda y derecha.
        """
        n_x = max(0, int(self.d.n_barras_intermedias_x))
        n_y = max(0, int(self.d.n_barras_intermedias_y))

        x_left = x0 + rec
        x_right = x0 + lbe - rec
        y_bot = rec
        y_top = bw - rec

        pts = []

        # 4 esquinas
        pts.extend([
            (x_left, y_bot),
            (x_right, y_bot),
            (x_left, y_top),
            (x_right, y_top),
        ])

        # Barras intermedias en caras X: inferior y superior
        for i in range(n_x):
            xi = x_left + (x_right - x_left) * (i + 1) / (n_x + 1)
            pts.append((xi, y_bot))
            pts.append((xi, y_top))

        # Barras intermedias en caras Y: izquierda y derecha
        for j in range(n_y):
            yj = y_bot + (y_top - y_bot) * (j + 1) / (n_y + 1)
            pts.append((x_left, yj))
            pts.append((x_right, yj))

        return pts

    def plot_corte_transversal(self, ax=None):
        """
        Vista en planta del muro completo:
        borde izquierdo + alma + borde derecho.
        """
        r = self.resolver()
        b = r["borde"]

        if ax is None:
            fig, ax = plt.subplots(figsize=(12, 4.8))
        else:
            fig = ax.figure

        lw = self.d.lw_m
        bw = self.d.bw_m

        def draw_hook_135(ax, x, y, orient="right_up", size=0.03, linewidth=1.8, color="black"):
            """
            Dibuja esquemáticamente un gancho de 135°.
            """
            kw = dict(linewidth=linewidth, color=color, solid_capstyle="round")
            if orient == "right_up":
                ax.plot([x, x + size], [y, y], **kw)
                ax.plot([x + size, x + 0.45*size], [y, y + 0.55*size], **kw)
            elif orient == "right_down":
                ax.plot([x, x + size], [y, y], **kw)
                ax.plot([x + size, x + 0.45*size], [y, y - 0.55*size], **kw)
            elif orient == "left_up":
                ax.plot([x, x - size], [y, y], **kw)
                ax.plot([x - size, x - 0.45*size], [y, y + 0.55*size], **kw)
            elif orient == "left_down":
                ax.plot([x, x - size], [y, y], **kw)
                ax.plot([x - size, x - 0.45*size], [y, y - 0.55*size], **kw)

        def draw_estribo_cerrado_135(ax, x0, y0, w, h, rec, linewidth=1.3, hook_size=0.026):
            """
            Dibuja el estribo cerrado del borde y muestra esquemáticamente ganchos de 135°
            en dos extremos del lazo.
            """
            xi = x0 + rec
            yi = y0 + rec
            wi = w - 2*rec
            hi = h - 2*rec

            ax.add_patch(Rectangle((xi, yi), wi, hi, fill=False, linewidth=linewidth, edgecolor="black"))

            # Ganchos de 135° esquemáticos en dos esquinas opuestas
            draw_hook_135(ax, xi, yi + hi, orient="right_down", size=hook_size, linewidth=linewidth+0.6, color="black")
            draw_hook_135(ax, xi + wi, yi, orient="left_up", size=hook_size, linewidth=linewidth+0.6, color="black")

        def draw_vincha_135(ax, x1, x2, y, linewidth=3.2, hook_size=0.034):
            """
            Dibuja una vincha paralela a lw con ganchos de 135° bien visibles en ambos extremos.
            """
            color = "black"
            margen = max(hook_size * 1.15, 0.032)

            # Cuerpo de la vincha ligeramente recortado para que los ganchos se distingan.
            ax.plot(
                [x1 + margen, x2 - margen],
                [y, y],
                linewidth=linewidth,
                color=color,
                solid_capstyle="round"
            )

            # Gancho izquierdo y derecho, saliendo hacia afuera de la vincha.
            draw_hook_135(
                ax, x1 + margen, y,
                orient="left_up",
                size=hook_size,
                linewidth=linewidth * 0.9,
                color=color
            )
            draw_hook_135(
                ax, x2 - margen, y,
                orient="right_down",
                size=hook_size,
                linewidth=linewidth * 0.9,
                color=color
            )

        ax.add_patch(Rectangle((0, 0), lw, bw, fill=False, linewidth=2.2))

        if b["requiere_borde_especial"]:
            lbe = b["lbe_cm"] / 100.0

            # Bordes integrados al muro
            ax.add_patch(Rectangle((0, 0), lbe, bw, alpha=0.18))
            ax.add_patch(Rectangle((lw - lbe, 0), lbe, bw, alpha=0.18))

            ax.plot([lbe, lbe], [0, bw], linestyle="--", linewidth=1.2)
            ax.plot([lw - lbe, lw - lbe], [0, bw], linestyle="--", linewidth=1.2)

            # Estribos esquemáticos en los bordes.
            # El estribo cerrado siempre se dibuja con ganchos de 135°.
            # Las vinchas interiores se dibujan paralelas a lw con ganchos de 135°.
            rec = 0.035
            n_vinchas = int(b.get("n_vinchas", 0))
            for x0 in [0, lw - lbe]:
                draw_estribo_cerrado_135(ax, x0, 0, lbe, bw, rec, linewidth=1.3, hook_size=0.022)

                if n_vinchas > 0:
                    for i in range(n_vinchas):
                        yv = rec + (bw - 2*rec) * (i + 1) / (n_vinchas + 1)
                        draw_vincha_135(
                            ax,
                            x0 + rec,
                            x0 + lbe - rec,
                            yv,
                            linewidth=2.8,
                            hook_size=0.020
                        )

            # Barras de borde según la distribución ingresada por el usuario:
            # 4 esquinas + intermedias en X + intermedias en Y.
            for x0 in [0, lw - lbe]:
                for (xb, yb) in self.coordenadas_barras_cabezal(x0, lbe, bw, rec=0.055):
                    ax.add_patch(Circle((xb, yb), 0.012, fill=False, linewidth=1.2))

            ax.text(lbe/2, bw + 0.055, "Borde\nconfinado", ha="center", va="bottom", fontsize=9)
            ax.text(lw - lbe/2, bw + 0.055, "Borde\nconfinado", ha="center", va="bottom", fontsize=9)
            ax.text(lw/2, bw + 0.055, "Alma del muro", ha="center", va="bottom", fontsize=10)

            ax.annotate("", xy=(0, -0.025), xytext=(lbe, -0.025), arrowprops=dict(arrowstyle="<->"))
            ax.text(lbe/2, -0.045, f"lbe = {lbe:.2f} m", ha="center", va="top", fontsize=9)

            ax.annotate("", xy=(lw - lbe, -0.025), xytext=(lw, -0.025), arrowprops=dict(arrowstyle="<->"))
            ax.text(lw - lbe/2, -0.045, f"lbe = {lbe:.2f} m", ha="center", va="top", fontsize=9)

            x_ini = lbe + self.d.s_web_cm/100.0/2
            x_fin = lw - lbe
            titulo = "Corte transversal del muro completo en planta\n(bordes/cabezales integrados al alma)"
            nota = f"Bordes: {r['flexocompresion_o_compresion']['armado_borde']} | {b['detalle_estribos']}\nAlma: Ø{self.d.db_web_mm:.0f} @ {self.d.s_web_cm:.0f} cm en dos caras"

        else:
            ax.text(lw/2, bw + 0.055, "Alma del muro\nsin borde especial", ha="center", va="bottom", fontsize=10)
            x_ini = self.d.s_web_cm/100.0/2
            x_fin = lw
            titulo = "Corte transversal del muro completo en planta\n(sin elementos especiales de borde)"
            nota = f"Alma: Ø{self.d.db_web_mm:.0f} @ {self.d.s_web_cm:.0f} cm en dos caras"

        # Acero distribuido del alma en dos caras
        x = x_ini
        while x < x_fin:
            ax.add_patch(Circle((x, 0.055), 0.009, fill=False, linewidth=0.9))
            ax.add_patch(Circle((x, bw - 0.055), 0.009, fill=False, linewidth=0.9))
            x += self.d.s_web_cm / 100.0

        ax.annotate("", xy=(0, -0.08), xytext=(lw, -0.08), arrowprops=dict(arrowstyle="<->"))
        ax.text(lw/2, -0.125, f"lw = {lw:.2f} m", ha="center", va="top")

        ax.annotate("", xy=(-0.08, 0), xytext=(-0.08, bw), arrowprops=dict(arrowstyle="<->"))
        ax.text(-0.12, bw/2, f"bw = {bw:.2f} m", rotation=90, ha="center", va="center")

        ax.set_title(titulo, fontsize=12)
        ax.text(lw/2, bw + 0.17, nota, ha="center", va="bottom", fontsize=10)

        ax.set_aspect("equal")
        ax.set_xlim(-0.18, lw + 0.18)
        ax.set_ylim(-0.16, bw + 0.32)
        ax.axis("off")
        return fig, ax


    def plot_figuras_separadas(self):
        """
        Genera dos figuras separadas para que no se reduzca demasiado el tamaño
        cuando el muro es muy largo.
        """
        fig1, ax1 = plt.subplots(figsize=(9, 7))
        self.plot_corte_longitudinal(ax=ax1)
        fig1.tight_layout()

        fig2, ax2 = plt.subplots(figsize=(13, 5))
        self.plot_corte_transversal(ax=ax2)
        fig2.tight_layout()

        return fig1, fig2

    def plot_resumen(self):
        fig, axs = plt.subplots(1, 2, figsize=(16, 6))
        self.plot_corte_longitudinal(ax=axs[0])
        self.plot_corte_transversal(ax=axs[1])
        fig.suptitle("Resumen gráfico del diseño del muro", fontsize=14)
        fig.tight_layout()
        return fig, axs


def fig_to_png_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=220, bbox_inches="tight")
    buf.seek(0)
    return buf.getvalue()


def dict_to_df(d):
    rows = []
    for k, v in d.items():
        if isinstance(v, float):
            rows.append({"Parámetro": k, "Valor": round(v, 4)})
        else:
            rows.append({"Parámetro": k, "Valor": v})
    return pd.DataFrame(rows)



def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_from_rows(doc, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for j, value in enumerate(rows[0]):
        set_cell_text(hdr[j], value, bold=True)

    for row in rows[1:]:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            set_cell_text(cells[j], value)

    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)

    return table


def add_equation_paragraph(doc, label, expression):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.bold = True
    p.add_run(" " + expression)


def generar_memoria_word(datos, muro, fig_long, fig_trans, info_memoria):
    """
    Genera una memoria de cálculo en Word con redacción dinámica,
    pasos explicativos, resultados y figuras.
    """
    r = muro.resolver()
    g = r["geometria"]
    c = r["cortante"]
    fx = r["flexocompresion_o_compresion"]
    b = r["borde"]
    web = r["acero_distribuido"]

    doc = Document()

    # Márgenes
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 3"].font.name = "Arial"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("MEMORIA DE CÁLCULO\nDISEÑO DE MURO ESTRUCTURAL")
    rt.bold = True
    rt.font.size = Pt(15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(info_memoria.get("nombre_proyecto", "Proyecto estructural")).bold = True

    meta_rows = [
        ["Elemento", info_memoria.get("nombre_muro", "Muro de corte")],
        ["Ejes", f"{info_memoria.get('eje_inicial', '#')} - {info_memoria.get('eje_final', '#')}"],
        ["Niveles", f"{info_memoria.get('nivel_inicial', '##')} a {info_memoria.get('nivel_final', '##')}"],
        ["Altura considerada", f"{datos.hw_m:.2f} m"],
        ["Longitud del muro", f"{datos.lw_m:.2f} m"],
        ["Espesor del muro", f"{datos.bw_m:.2f} m"],
    ]
    if info_memoria.get("autor_memoria", "").strip():
        meta_rows.append(["Elaborado por", info_memoria.get("autor_memoria")])
    add_table_from_rows(doc, [["Dato", "Descripción"]] + meta_rows, col_widths=[1.7, 4.8])

    doc.add_heading("1. Alcance del cálculo", level=1)
    doc.add_paragraph(
        f"En el presente capítulo se muestra el cálculo para obtener el espesor y el armado del "
        f"{info_memoria.get('nombre_muro', 'muro de corte')}, ubicado entre los ejes "
        f"{info_memoria.get('eje_inicial', '#')}-{info_memoria.get('eje_final', '#')}. "
        f"El muro tiene una longitud Lw = {datos.lw_m:.2f} m y un espesor bw = {datos.bw_m:.2f} m. "
        f"El elemento se desarrolla desde el nivel {info_memoria.get('nivel_inicial', '##')} hasta el nivel "
        f"{info_memoria.get('nivel_final', '##')}, por lo que se considera una altura total hw = {datos.hw_m:.2f} m."
    )
    doc.add_paragraph(
        "El diseño se realiza considerando el comportamiento del muro en su plano, revisando la demanda de "
        "cortante amplificada, la flexo-compresión, el acero distribuido del alma y el detallamiento de los "
        "elementos especiales de borde o cabezales."
    )

    doc.add_heading("2. Datos de entrada", level=1)
    datos_rows = [
        ["Parámetro", "Valor", "Descripción"],
        ["Lw", f"{datos.lw_m:.2f} m", "Longitud total del muro"],
        ["bw", f"{datos.bw_m:.2f} m", "Espesor del muro"],
        ["hw", f"{datos.hw_m:.2f} m", "Altura total considerada"],
        ["N", f"{datos.n_pisos}", "Número de pisos"],
        ["f'c", f"{datos.fc_kgcm2:.0f} kg/cm²", "Resistencia a compresión del hormigón"],
        ["fy", f"{datos.fy_kgcm2:.0f} kg/cm²", "Fluencia del acero longitudinal"],
        ["fyt", f"{datos.fyt_kgcm2:.0f} kg/cm²", "Fluencia del acero transversal"],
        ["Pu", f"{datos.Pu_tonf:.2f} tonf", "Carga axial última"],
        ["Vu", f"{datos.Vu_tonf:.2f} tonf", "Cortante de análisis en el plano del muro"],
        ["Mu", f"{datos.Mu_plano_tonfm:.2f} t·m", "Momento último en el plano del muro"],
    ]
    add_table_from_rows(doc, datos_rows, col_widths=[1.1, 1.8, 4.2])

    doc.add_heading("3. Revisión por cortante en el plano", level=1)
    doc.add_paragraph(
        "El cortante ingresado corresponde al cortante que actúa en el plano del muro. "
        "Para la revisión se amplifica la fuerza de corte mediante el factor dinámico asociado al número de pisos "
        "y el factor de sobrerresistencia."
    )
    add_equation_paragraph(doc, "Factor dinámico:", f"ω = 1.3 + N/30 = 1.3 + {datos.n_pisos}/30 = {c['omega']:.2f}")
    add_equation_paragraph(doc, "Cortante amplificado:", f"Vu' = ω·φo·Vu = {c['omega']:.2f}·{datos.phi_o:.2f}·{datos.Vu_tonf:.2f} = {c['Vu_amplificado_tonf']:.2f} tonf")
    add_equation_paragraph(doc, "Resistencia a cortante:", f"φVn = φ·1.6·√f'c·bw·lw = {c['phiVn_tonf']:.2f} tonf")
    doc.add_paragraph(
        f"Con la geometría adoptada se obtiene φVn = {c['phiVn_tonf']:.2f} tonf, "
        f"mientras que el cortante amplificado es Vu' = {c['Vu_amplificado_tonf']:.2f} tonf. "
        f"Por lo tanto, la revisión por cortante en el plano del muro "
        f"{'cumple' if c['cumple_cortante'] else 'no cumple'}."
    )

    cort_rows = [
        ["Resultado", "Valor"],
        ["Vu amplificado", f"{c['Vu_amplificado_tonf']:.2f} tonf"],
        ["φVn", f"{c['phiVn_tonf']:.2f} tonf"],
        ["Lw requerido", f"{c['lw_req_cm']:.1f} cm"],
        ["bw requerido", f"{c['bw_req_cm']:.1f} cm"],
        ["Veredicto", "CUMPLE" if c["cumple_cortante"] else "NO CUMPLE"],
    ]
    add_table_from_rows(doc, cort_rows, col_widths=[2.4, 2.4])

    doc.add_heading("4. Revisión por flexo-compresión en el plano", level=1)
    if fx["modo"] == "COMPRESION_AXIAL_PREDOMINANTE":
        doc.add_paragraph(
            "Debido a que el momento en el plano ingresado es igual a cero, el muro se revisa como un elemento "
            "con compresión axial predominante."
        )
        add_equation_paragraph(doc, "Resistencia axial:", f"φPn = {fx['phiPn_tonf']:.2f} tonf")
        doc.add_paragraph(
            f"La resistencia axial de diseño φPn = {fx['phiPn_tonf']:.2f} tonf se compara con "
            f"Pu = {datos.Pu_tonf:.2f} tonf. La revisión axial "
            f"{'cumple' if fx['cumple_axial'] else 'no cumple'}."
        )
    else:
        doc.add_paragraph(
            "Para el diseño a flexo-compresión se calcula el momento nominal requerido a partir del momento último "
            "y del factor de reducción de resistencia a flexión. Luego se considera el aporte del acero distribuido "
            "vertical real colocado en el alma del muro y se obtiene el acero adicional requerido en los cabezales."
        )
        add_equation_paragraph(doc, "Momento nominal:", f"Mn = Mu/φ = {datos.Mu_plano_tonfm:.2f}/{datos.phi_flexion:.2f} = {fx['Mn_tonfm']:.2f} t·m")
        add_equation_paragraph(doc, "Brazo de la carga axial:", f"xp = {datos.xp_factor_lw:.3f}·Lw = {fx['xp_m']:.3f} m")
        add_equation_paragraph(doc, "Acero vertical del alma:", f"As1 = As/m · 0.60Lw = {fx['As_vertical_alma_por_metro_cm2_m']:.2f}·{fx['longitud_alma_efectiva_m']:.2f} = {fx['As1_cm2']:.2f} cm²")
        add_equation_paragraph(doc, "Tracción aportada por el alma:", f"Ts1 = As1·fy = {fx['Ts1_tonf']:.2f} tonf")
        add_equation_paragraph(doc, "Tracción restante:", f"Ts2 = {fx['Ts2_tonf']:.2f} tonf")
        add_equation_paragraph(doc, "Acero adicional de borde:", f"As2 = Ts2/fy = {fx['As2_req_cm2']:.2f} cm²")

        if fx["requiere_acero_adicional_borde_por_flexion"]:
            doc.add_paragraph(
                f"El acero adicional requerido por flexión en cada cabezal es As2 = {fx['As2_req_cm2']:.2f} cm². "
                f"Con barras Ø{datos.db_borde_mm:.0f} mm, el armado mínimo equivalente es "
                f"{fx['armado_borde_requerido_minimo']}. El usuario ha colocado {fx['armado_borde']}, "
                f"con un área As = {fx['As2_colocado_cm2']:.2f} cm², por lo que el acero de borde "
                f"{'cumple' if fx['cumple_acero_borde'] else 'no cumple'}."
            )
        else:
            doc.add_paragraph(
                "El equilibrio indica que no se requiere acero adicional de borde por flexión, ya que el acero vertical "
                "distribuido del alma cubre la tracción necesaria. Sin embargo, el cabezal puede seguir siendo necesario "
                "por criterios de confinamiento del borde comprimido."
            )

        flex_rows = [
            ["Resultado", "Valor"],
            ["Mn", f"{fx['Mn_tonfm']:.2f} t·m"],
            ["xp", f"{fx['xp_m']:.3f} m"],
            ["As1 real aportante", f"{fx['As1_cm2']:.2f} cm²"],
            ["Ts1", f"{fx['Ts1_tonf']:.2f} tonf"],
            ["Ts2", f"{fx['Ts2_tonf']:.2f} tonf"],
            ["As2 requerido", f"{fx['As2_req_cm2']:.2f} cm²"],
            ["As2 colocado", f"{fx['As2_colocado_cm2']:.2f} cm²"],
            ["Veredicto", "CUMPLE" if fx["cumple_acero_borde"] else "NO CUMPLE"],
        ]
        add_table_from_rows(doc, flex_rows, col_widths=[2.6, 2.4])

    doc.add_heading("5. Elementos especiales de borde o cabezales", level=1)
    if b["requiere_borde_especial"]:
        doc.add_paragraph(
            "Se revisa la necesidad y el detallamiento de los elementos especiales de borde. "
            "El desplazamiento máximo de techo se utiliza para calcular la relación δu/hw y, con ello, "
            "la profundidad del eje neutro asociada al criterio de borde."
        )
        add_equation_paragraph(doc, "Relación de desplazamiento:", f"δu/hw = {b['desplazamiento_techo_cm']:.2f}/{datos.hw_m*100:.1f} = {b['delta_u_sobre_hw']:.6f}")
        add_equation_paragraph(doc, "Profundidad calculada:", f"c = {b['c_cm']:.1f} cm; c/Lw = {b['c_sobre_lw']:.3f}")
        add_equation_paragraph(doc, "Longitud del elemento de borde:", f"lbe = max(c - 0.10Lw, c/2) = {b['lbe_cm']:.1f} cm")
        doc.add_paragraph(
            f"Con estos valores se adopta una longitud de elemento de borde lbe = {b['lbe_cm']:.1f} cm en cada extremo del muro. "
            f"El espaciamiento de estribos resulta s = {b['s_estribos_cm']:.1f} cm, controlado por {b['controla_s']}."
        )
        borde_rows = [
            ["Resultado", "Valor"],
            ["Pu/(Ag·f'c)", f"{b['Pu_sobre_Agfc']:.3f}"],
            ["δu", f"{b['desplazamiento_techo_cm']:.2f} cm"],
            ["δu/hw", f"{b['delta_u_sobre_hw']:.6f}"],
            ["c", f"{b['c_cm']:.1f} cm"],
            ["c/Lw", f"{b['c_sobre_lw']:.3f}"],
            ["lbe", f"{b['lbe_cm']:.1f} cm"],
            ["ldh", f"{b['ldh_cm']:.1f} cm"],
            ["hx usado", f"{b['hx_mm']:.0f} mm"],
            ["s estribos", f"{b['s_estribos_cm']:.1f} cm"],
            ["Ash1", f"{b['Ash1_req_cm2']:.2f} cm² → {b['n_ramas_Ash1']} ramas"],
            ["Ash2", f"{b['Ash2_req_cm2']:.2f} cm² → {b['n_ramas_Ash2']} ramas"],
            ["Detalle", b["detalle_estribos"]],
        ]
        add_table_from_rows(doc, borde_rows, col_widths=[2.5, 3.2])
    else:
        doc.add_paragraph("No se requieren elementos especiales de borde debido a que no existe momento en el plano del muro.")

    doc.add_heading("6. Acero distribuido del alma", level=1)
    doc.add_paragraph(
        f"El alma del muro se arma con {web['detalle']}. Este acero se dispone en dos caras del muro y se verifica "
        "frente a las cuantías mínimas vertical y horizontal."
    )
    web_rows = [
        ["Resultado", "Valor"],
        ["Armado adoptado", web["detalle"]],
        ["As por metro", f"{web['As_por_metro_cm2_m']:.2f} cm²/m"],
        ["Cuantía colocada", f"{web['rho_colocada']:.5f}"],
        ["Cuantía vertical mínima", "CUMPLE" if web["cumple_rho_vertical"] else "NO CUMPLE"],
        ["Cuantía horizontal mínima", "CUMPLE" if web["cumple_rho_horizontal"] else "NO CUMPLE"],
    ]
    add_table_from_rows(doc, web_rows, col_widths=[2.6, 3.4])

    doc.add_heading("7. Esquemas de armado", level=1)
    doc.add_paragraph(
        "A continuación se presentan los esquemas gráficos generados con los datos ingresados. "
        "Las figuras son referenciales para la memoria y muestran la ubicación de los elementos de borde, "
        "el acero distribuido del alma, estribos, vinchas y ganchos."
    )

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        long_path = tmpdir / "corte_longitudinal.png"
        trans_path = tmpdir / "corte_transversal.png"
        fig_long.savefig(long_path, dpi=220, bbox_inches="tight")
        fig_trans.savefig(trans_path, dpi=220, bbox_inches="tight")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Figura 1. Corte longitudinal del muro.").bold = True
        doc.add_picture(str(long_path), width=Inches(5.8))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Figura 2. Corte transversal en planta del muro.").bold = True
        doc.add_picture(str(trans_path), width=Inches(6.5))

    doc.add_heading("8. Conclusión", level=1)
    conclusiones = []
    conclusiones.append("La revisión por cortante en el plano del muro " + ("cumple." if c["cumple_cortante"] else "no cumple."))
    if fx["modo"] == "FLEXO_COMPRESION_EN_EL_PLANO":
        conclusiones.append("La revisión del acero adicional de borde por flexión " + ("cumple." if fx["cumple_acero_borde"] else "no cumple."))
    else:
        conclusiones.append("La revisión por compresión axial predominante " + ("cumple." if fx["cumple_axial"] else "no cumple."))
    conclusiones.append("El acero distribuido del alma " + ("cumple las cuantías mínimas." if web["cumple_rho_vertical"] and web["cumple_rho_horizontal"] else "no cumple alguna cuantía mínima."))
    for item in conclusiones:
        doc.add_paragraph(item, style=None)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def main():
    st.set_page_config(
        page_title="Diseño de muro estructural",
        page_icon="🏗️",
        layout="wide",
    )

    st.title("🏗️ Diseño de muro estructural de hormigón armado")
    st.caption("Flujo único: ingresa datos → memoria de cálculo → veredictos → gráficos.")

    with st.sidebar:
        st.header("Datos de entrada")

        st.subheader("1. Geometría")
        lw_m = st.number_input("Longitud del muro, lw [m]", min_value=0.10, value=5.00, step=0.10, format="%.2f")
        bw_m = st.number_input("Espesor del muro, bw [m]", min_value=0.05, value=0.30, step=0.01, format="%.2f")
        hw_m = st.number_input("Altura total del muro, hw [m]", min_value=0.50, value=33.66, step=0.10, format="%.2f")
        n_pisos = st.number_input("Número de pisos, N", min_value=1, value=11, step=1)

        with st.expander("Identificación para memoria Word"):
            nombre_proyecto = st.text_input("Nombre del proyecto", value="Proyecto estructural")
            nombre_muro = st.text_input("Identificación del muro", value="Muro de corte M-01")
            eje_inicial = st.text_input("Eje inicial", value="A")
            eje_final = st.text_input("Eje final", value="B")
            nivel_inicial = st.text_input("Nivel inicial", value="N+0.00")
            nivel_final = st.text_input("Nivel final", value="N+15.00")
            autor_memoria = st.text_input("Elaborado por", value="")

        st.subheader("2. Materiales")
        fc_kgcm2 = st.number_input("f'c [kg/cm²]", min_value=100.0, value=280.0, step=10.0, format="%.0f")
        fy_kgcm2 = st.number_input("fy longitudinal [kg/cm²]", min_value=2000.0, value=4200.0, step=100.0, format="%.0f")
        fyt_kgcm2 = st.number_input("fyt transversal [kg/cm²]", min_value=2000.0, value=4200.0, step=100.0, format="%.0f")

        st.subheader("3. Solicitaciones")
        Pu_tonf = st.number_input("Pu [tonf]", min_value=0.0, value=352.0, step=1.0, format="%.2f")
        Vu_tonf = st.number_input("Vu de análisis [tonf]", min_value=0.0, value=60.0, step=1.0, format="%.2f")
        Mu_plano_tonfm = st.number_input("Mu en el plano [t·m]", min_value=0.0, value=1875.0, step=1.0, format="%.2f")

        st.subheader("4. Factores")
        phi_v = st.number_input("φ cortante", min_value=0.10, max_value=1.00, value=0.75, step=0.01, format="%.2f")
        phi_flexion = st.number_input("φ flexión", min_value=0.10, max_value=1.00, value=0.90, step=0.01, format="%.2f")
        phi_o = st.number_input("φo sobrerresistencia", min_value=0.10, value=1.40, step=0.05, format="%.2f")

        with st.expander("Cuantías y parámetros avanzados"):
            rho_vertical = st.number_input("ρ vertical mínima", min_value=0.0, value=0.0025, step=0.0001, format="%.5f")
            rho_horizontal = st.number_input("ρ horizontal mínima", min_value=0.0, value=0.0025, step=0.0001, format="%.5f")
            desplazamiento_techo_cm = st.number_input(
                "Desplazamiento máximo de techo δu [cm]",
                min_value=0.01,
                value=14.96,
                step=0.10,
                format="%.2f",
                help="Valor tomado del análisis estructural. El programa calcula internamente δu/hw = δu / hw."
            )
            xp_factor_lw = st.number_input("xp/lw", min_value=0.0, max_value=1.0, value=0.095, step=0.005, format="%.3f")
            bc_ash_1_cm = st.number_input("bc Ash1 [cm]", min_value=1.0, value=40.0, step=1.0, format="%.1f")
            bc_ash_2_cm = st.number_input("bc Ash2 [cm]", min_value=1.0, value=26.7, step=0.1, format="%.1f")
            modo_hx_ui = st.radio(
                "Cálculo de hx",
                ["Calcular hx automáticamente", "Ingresar hx manualmente"],
                index=0,
                help="Usa manual cuando tu detalle real tenga vinchas/ramales adicionales y quieras ingresar el hx real medido entre barras soportadas."
            )

            recubrimiento_libre_cm = st.number_input(
                "Recubrimiento libre del cabezal/muro [cm]",
                min_value=0.0,
                value=4.0,
                step=0.5,
                format="%.1f",
                help="Recubrimiento puro hasta el estribo. El programa calcula internamente la distancia al centro de la barra extrema: recubrimiento + Øestribo + Ølongitudinal/2."
            )

            hx_manual_mm = 200.0
            if modo_hx_ui == "Ingresar hx manualmente":
                hx_manual_mm = st.number_input(
                    "hx manual [mm]",
                    min_value=1.0,
                    value=200.0,
                    step=5.0,
                    format="%.0f",
                    help="Ingresa el hx real de tu detalle: mayor separación centro a centro entre barras longitudinales soportadas lateralmente por estribos/vinchas/ganchos."
                )

            modo_hx = "manual" if modo_hx_ui == "Ingresar hx manualmente" else "auto"

        st.subheader("5. Cabezal / elemento de borde")
        db_borde_mm = st.number_input("Ø longitudinal de cabezal [mm]", min_value=6.0, value=20.0, step=1.0, format="%.0f")
        n_barras_esquinas = st.number_input("Barras en esquinas", min_value=4, value=4, step=1)
        n_barras_intermedias_x = st.number_input("Intermedias en X por cara", min_value=0, value=2, step=1)
        n_barras_intermedias_y = st.number_input("Intermedias en Y por cara", min_value=0, value=2, step=1)
        db_estribo_mm = st.number_input("Ø estribo/vincha [mm]", min_value=6.0, value=10.0, step=1.0, format="%.0f")

        st.subheader("6. Alma del muro")
        db_web_mm = st.number_input("Ø acero distribuido del alma [mm]", min_value=6.0, value=20.0, step=1.0, format="%.0f")
        s_web_cm = st.number_input("Separación alma [cm]", min_value=5.0, value=25.0, step=1.0, format="%.0f")

    datos = DatosMuro(
        lw_m=lw_m,
        bw_m=bw_m,
        hw_m=hw_m,
        n_pisos=int(n_pisos),
        fc_kgcm2=fc_kgcm2,
        fy_kgcm2=fy_kgcm2,
        fyt_kgcm2=fyt_kgcm2,
        Pu_tonf=Pu_tonf,
        Vu_tonf=Vu_tonf,
        Mu_plano_tonfm=Mu_plano_tonfm,
        phi_v=phi_v,
        phi_flexion=phi_flexion,
        phi_o=phi_o,
        rho_vertical=rho_vertical,
        rho_horizontal=rho_horizontal,
        desplazamiento_techo_cm=desplazamiento_techo_cm,
        db_borde_mm=db_borde_mm,
        n_barras_esquinas=int(n_barras_esquinas),
        n_barras_intermedias_x=int(n_barras_intermedias_x),
        n_barras_intermedias_y=int(n_barras_intermedias_y),
        db_estribo_mm=db_estribo_mm,
        db_web_mm=db_web_mm,
        s_web_cm=s_web_cm,
        xp_factor_lw=xp_factor_lw,
        bc_ash_1_cm=bc_ash_1_cm,
        bc_ash_2_cm=bc_ash_2_cm,
        recubrimiento_libre_cm=recubrimiento_libre_cm,
        modo_hx=modo_hx,
        hx_manual_mm=hx_manual_mm,
    )

    info_memoria = {
        "nombre_proyecto": nombre_proyecto,
        "nombre_muro": nombre_muro,
        "eje_inicial": eje_inicial,
        "eje_final": eje_final,
        "nivel_inicial": nivel_inicial,
        "nivel_final": nivel_final,
        "autor_memoria": autor_memoria,
    }

    muro = DisenoMuro(datos)
    resultados = muro.resolver()
    fx = resultados["flexocompresion_o_compresion"]
    cortante = resultados["cortante"]
    borde = resultados["borde"]
    web = resultados["acero_distribuido"]

    st.subheader("Resumen rápido")
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Modo", resultados["modo"].replace("_", " "))
    col2.metric("Vu' [tonf]", f"{cortante['Vu_amplificado_tonf']:.2f}")
    col3.metric("Cortante", "CUMPLE" if cortante["cumple_cortante"] else "NO CUMPLE")
    if fx["modo"] == "FLEXO_COMPRESION_EN_EL_PLANO":
        col4.metric("Acero borde", "CUMPLE" if fx["cumple_acero_borde"] else "NO CUMPLE")
    else:
        col4.metric("Axial", "CUMPLE" if fx["cumple_axial"] else "NO CUMPLE")

    tab_memoria, tab_graficos, tab_tablas, tab_validacion = st.tabs(["Memoria", "Gráficos", "Tablas", "Validación"] )

    with tab_memoria:
        memoria = muro.memoria_markdown()
        st.markdown(memoria)

        # Figuras para insertar en la memoria Word.
        fig_long_docx, ax_long_docx = plt.subplots(figsize=(9, 7))
        muro.plot_corte_longitudinal(ax=ax_long_docx)
        fig_trans_docx, ax_trans_docx = plt.subplots(figsize=(13, 5))
        muro.plot_corte_transversal(ax=ax_trans_docx)

        memoria_word = generar_memoria_word(datos, muro, fig_long_docx, fig_trans_docx, info_memoria)

        st.download_button(
            "Descargar memoria Word (.docx)",
            data=memoria_word,
            file_name="memoria_calculo_muro.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        st.download_button(
            "Descargar memoria Markdown",
            data=memoria.encode("utf-8"),
            file_name="memoria_diseno_muro.md",
            mime="text/markdown",
        )

    with tab_graficos:
        st.markdown("### Corte longitudinal")
        fig_long, ax_long = plt.subplots(figsize=(9, 7))
        muro.plot_corte_longitudinal(ax=ax_long)
        st.pyplot(fig_long, clear_figure=False)
        png_long = fig_to_png_bytes(fig_long)
        st.download_button(
            "Descargar corte longitudinal PNG",
            data=png_long,
            file_name="muro_corte_longitudinal.png",
            mime="image/png",
        )

        st.markdown("### Corte transversal")
        fig_trans, ax_trans = plt.subplots(figsize=(13, 5))
        muro.plot_corte_transversal(ax=ax_trans)
        st.pyplot(fig_trans, clear_figure=False)
        png_trans = fig_to_png_bytes(fig_trans)
        st.download_button(
            "Descargar corte transversal PNG",
            data=png_trans,
            file_name="muro_corte_transversal.png",
            mime="image/png",
        )

    with tab_tablas:
        st.markdown("### Geometría")
        st.dataframe(dict_to_df(resultados["geometria"]), use_container_width=True)
        st.markdown("### Cortante")
        st.dataframe(dict_to_df(cortante), use_container_width=True)
        st.markdown("### Flexo-compresión o compresión")
        st.dataframe(dict_to_df(fx), use_container_width=True)
        st.markdown("### Elementos de borde")
        st.dataframe(dict_to_df(borde), use_container_width=True)
        st.markdown("### Acero distribuido del alma")
        st.dataframe(dict_to_df(web), use_container_width=True)

    with tab_validacion:
        val = muro.validar_ejemplo_documento()
        if val["aplica_validacion"]:
            st.info("La validación exacta puede cambiar si modificas el acero real del alma, porque ahora As1 depende de db_web_mm y s_web_cm.")
            st.write("Resultado:", "OK" if val["ok_validacion"] else "REVISAR")
            st.dataframe(pd.DataFrame(val["comparacion"]).T, use_container_width=True)
        else:
            st.info(val["mensaje"])

    plt.close('all')


if __name__ == "__main__":
    main()
